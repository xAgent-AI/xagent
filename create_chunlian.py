from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 创建Word文档
doc = Document()

# 设置页面为纵向（A4）
section = doc.sections[0]
section.page_height = Inches(11.69)
section.page_width = Inches(8.27)

# 设置边距
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

# 添加标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('2025新春对联')
run.font.size = Pt(28)
run.font.name = '宋体'
run.font.color.rgb = RGBColor(200, 0, 0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 添加空行
doc.add_paragraph()

# 上联
shanglian = doc.add_paragraph()
shanglian.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = shanglian.add_run('上联：春回大地千山秀')
run.font.size = Pt(24)
run.font.name = '楷体'
run.font.color.rgb = RGBColor(180, 0, 0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# 下联
xialian = doc.add_paragraph()
xialian.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = xialian.add_run('下联：日暖神州万物荣')
run.font.size = Pt(24)
run.font.name = '楷体'
run.font.color.rgb = RGBColor(180, 0, 0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# 横批
hengpi = doc.add_paragraph()
hengpi.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = hengpi.add_run('横批：万象更新')
run.font.size = Pt(26)
run.font.name = '黑体'
run.font.color.rgb = RGBColor(220, 0, 0)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 添加分隔线
doc.add_paragraph()

# 第二副对联
doc.add_paragraph().add_run('=' * 50).font.size = Pt(12)

# 上联
shanglian2 = doc.add_paragraph()
shanglian2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = shanglian2.add_run('上联：瑞雪迎春铺锦绣')
run.font.size = Pt(24)
run.font.name = '楷体'
run.font.color.rgb = RGBColor(180, 0, 0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# 下联
xialian2 = doc.add_paragraph()
xialian2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = xialian2.add_run('下联：红梅报岁展宏图')
run.font.size = Pt(24)
run.font.name = '楷体'
run.font.color.rgb = RGBColor(180, 0, 0)
run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

# 横批
hengpi2 = doc.add_paragraph()
hengpi2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = hengpi2.add_run('横批：新春大吉')
run.font.size = Pt(26)
run.font.name = '黑体'
run.font.color.rgb = RGBColor(220, 0, 0)
run.bold = True
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 保存文档
doc.save('春联.docx')
print('春联文档已成功生成：春联.docx')
